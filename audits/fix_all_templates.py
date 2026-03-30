#!/usr/bin/env python3
"""
Fix all templates:
1. Swap cover logo to Maker Logo - Official.png (no white sliver)
2. Remove orange container border around cover logo
3. Rebuild footer with logo + text + QR code properly embedded
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree
import os
import shutil

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")
LOGO_CLEAN = "/Users/guillermo/Desktop/Rex/Rex/projects/Maker Web Studios/public/logos/Maker Logo - Official.png"
QR_PATH = os.path.join(TEMPLATE_DIR, "mws_qr_code.png")

ns_w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
ns_r = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
ns_wp = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
ns_a = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
ns_pic = '{http://schemas.openxmlformats.org/drawingml/2006/picture}'


def swap_cover_logo(doc):
    """Replace the cover page logo image with the clean Official version."""
    if not doc.tables:
        return False

    table = doc.tables[0]
    # Find inline images in the cover table
    for inline in table._element.findall(f'.//{ns_wp}inline'):
        # Get the blip element (actual image reference)
        blip = inline.find(f'.//{ns_a}blip')
        if blip is not None:
            old_rid = blip.get(qn('r:embed'))
            if old_rid:
                # Get the document part and add the new image
                part = doc.part
                # Add the clean logo as a new image
                with open(LOGO_CLEAN, 'rb') as f:
                    image_data = f.read()

                from docx.opc.part import Part
                from docx.image.image import Image

                image = Image.from_blob(image_data)

                # Add image to package
                img_part = part.package.get_or_add_image_part(LOGO_CLEAN)
                new_rid = part.relate_to(img_part, RT.IMAGE)

                # Update blip to reference new image
                blip.set(qn('r:embed'), new_rid)

                # Remove any outline/border on the shape
                sp_pr = inline.find(f'.//{ns_a}spPr')
                if sp_pr is not None:
                    # Remove any line elements (borders)
                    for ln in sp_pr.findall(f'{ns_a}ln'):
                        sp_pr.remove(ln)
                    # Add explicit "no line"
                    ln = etree.SubElement(sp_pr, qn('a:ln'))
                    no_fill = etree.SubElement(ln, qn('a:noFill'))

                return True
    return False


def swap_header_logo(doc):
    """Replace header logo with clean version too."""
    for section in doc.sections:
        if not section.header:
            continue
        header_part = section.header.part
        for p in section.header.paragraphs:
            for run in p.runs:
                for inline in run._element.findall(f'.//{ns_wp}inline'):
                    blip = inline.find(f'.//{ns_a}blip')
                    if blip is not None:
                        old_rid = blip.get(qn('r:embed'))
                        if old_rid:
                            img_part = doc.part.package.get_or_add_image_part(LOGO_CLEAN)
                            new_rid = header_part.relate_to(img_part, RT.IMAGE)
                            blip.set(qn('r:embed'), new_rid)

                            # Remove any border
                            sp_pr = inline.find(f'.//{ns_a}spPr')
                            if sp_pr is not None:
                                for ln in sp_pr.findall(f'{ns_a}ln'):
                                    sp_pr.remove(ln)
                                ln = etree.SubElement(sp_pr, qn('a:ln'))
                                etree.SubElement(ln, qn('a:noFill'))


def rebuild_footer(doc):
    """
    Rebuild footer with: logo (left) | text (center) | QR code (right)
    Images are added directly to the footer part for proper embedding.
    """
    for section in doc.sections:
        footer = section.footer
        if not footer:
            continue

        footer_part = footer.part

        # Clear footer
        for elem in list(footer._element):
            footer._element.remove(elem)

        # Add logo image to footer part
        logo_img_part = doc.part.package.get_or_add_image_part(LOGO_CLEAN)
        logo_rid = footer_part.relate_to(logo_img_part, RT.IMAGE)

        # Add QR image to footer part
        qr_img_part = doc.part.package.get_or_add_image_part(QR_PATH)
        qr_rid = footer_part.relate_to(qr_img_part, RT.IMAGE)

        # Build the footer table XML with proper image references
        logo_cx = int(0.35 * 914400)  # 0.35 inches in EMUs
        logo_cy = int(0.35 * 914400)
        qr_cx = int(0.35 * 914400)
        qr_cy = int(0.35 * 914400)

        footer_xml = f'''<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
               xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
               xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
          <w:tblPr>
            <w:tblW w:w="5000" w:type="pct"/>
            <w:tblBorders>
              <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
              <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
              <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
              <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
              <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
              <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            </w:tblBorders>
          </w:tblPr>
          <w:tr>
            <w:tc>
              <w:tcPr>
                <w:tcW w:w="1000" w:type="pct"/>
                <w:vAlign w:val="center"/>
              </w:tcPr>
              <w:p>
                <w:pPr><w:jc w:val="left"/></w:pPr>
                <w:r>
                  <w:drawing>
                    <wp:inline distT="0" distB="0" distL="0" distR="0">
                      <wp:extent cx="{logo_cx}" cy="{logo_cy}"/>
                      <wp:docPr id="100" name="Footer Logo"/>
                      <a:graphic>
                        <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                          <pic:pic>
                            <pic:nvPicPr>
                              <pic:cNvPr id="100" name="footer_logo.png"/>
                              <pic:cNvPicPr/>
                            </pic:nvPicPr>
                            <pic:blipFill>
                              <a:blip r:embed="{logo_rid}"/>
                              <a:stretch><a:fillRect/></a:stretch>
                            </pic:blipFill>
                            <pic:spPr>
                              <a:xfrm>
                                <a:off x="0" y="0"/>
                                <a:ext cx="{logo_cx}" cy="{logo_cy}"/>
                              </a:xfrm>
                              <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                              <a:ln><a:noFill/></a:ln>
                            </pic:spPr>
                          </pic:pic>
                        </a:graphicData>
                      </a:graphic>
                    </wp:inline>
                  </w:drawing>
                </w:r>
              </w:p>
            </w:tc>
            <w:tc>
              <w:tcPr>
                <w:tcW w:w="3000" w:type="pct"/>
                <w:vAlign w:val="center"/>
              </w:tcPr>
              <w:p>
                <w:pPr><w:jc w:val="left"/></w:pPr>
                <w:r>
                  <w:rPr>
                    <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>
                    <w:sz w:val="14"/>
                    <w:szCs w:val="14"/>
                    <w:color w:val="7E8691"/>
                  </w:rPr>
                  <w:t xml:space="preserve">www.makerwebstudios.com  |  CONFIDENTIAL — PREPARED EXCLUSIVELY FOR [CLIENT NAME]</w:t>
                </w:r>
              </w:p>
            </w:tc>
            <w:tc>
              <w:tcPr>
                <w:tcW w:w="1000" w:type="pct"/>
                <w:vAlign w:val="center"/>
              </w:tcPr>
              <w:p>
                <w:pPr><w:jc w:val="right"/></w:pPr>
                <w:r>
                  <w:drawing>
                    <wp:inline distT="0" distB="0" distL="0" distR="0">
                      <wp:extent cx="{qr_cx}" cy="{qr_cy}"/>
                      <wp:docPr id="101" name="QR Code"/>
                      <a:graphic>
                        <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                          <pic:pic>
                            <pic:nvPicPr>
                              <pic:cNvPr id="101" name="qr_code.png"/>
                              <pic:cNvPicPr/>
                            </pic:nvPicPr>
                            <pic:blipFill>
                              <a:blip r:embed="{qr_rid}"/>
                              <a:stretch><a:fillRect/></a:stretch>
                            </pic:blipFill>
                            <pic:spPr>
                              <a:xfrm>
                                <a:off x="0" y="0"/>
                                <a:ext cx="{qr_cx}" cy="{qr_cy}"/>
                              </a:xfrm>
                              <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                              <a:ln><a:noFill/></a:ln>
                            </pic:spPr>
                          </pic:pic>
                        </a:graphicData>
                      </a:graphic>
                    </wp:inline>
                  </w:drawing>
                </w:r>
              </w:p>
            </w:tc>
          </w:tr>
        </w:tbl>'''

        tbl_elem = etree.fromstring(footer_xml)
        footer._element.append(tbl_elem)


# Process all templates
for fname in sorted(os.listdir(TEMPLATE_DIR)):
    if not fname.endswith('.docx'):
        continue
    fpath = os.path.join(TEMPLATE_DIR, fname)
    doc = Document(fpath)

    # Fix 1: Swap cover logo
    swapped_cover = swap_cover_logo(doc)

    # Fix 2: Swap header logo
    swap_header_logo(doc)

    # Fix 3: Rebuild footer with proper image embedding
    rebuild_footer(doc)

    doc.save(fpath)
    print(f"  OK: {fname} (cover logo: {'swapped' if swapped_cover else 'no image found'})")

print("\nAll templates updated.")
