#!/usr/bin/env python3
"""
Fix double logo on cover pages and rebuild footers across all templates.
1. Remove redundant "MAKER" text from cover page tables
2. Rebuild footer: logo (left) | text (center) | QR code (right)
"""
from docx import Document
from docx.shared import Inches, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
import os
import copy

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")
LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "public", "logos", "Maker Logo - 300 x 300 px - Official.png")
QR_PATH = os.path.join(TEMPLATE_DIR, "mws_qr_code.png")

# Fallback logo path
if not os.path.exists(LOGO_PATH):
    LOGO_PATH = os.path.join(TEMPLATE_DIR, "..", "..", "public", "logos", "Maker Logo - 300 x 300 px - Official.png")

ns_w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def remove_maker_text_from_cover(doc):
    """
    Remove the redundant 'MAKER' text element from the cover page table.
    The logo image already displays 'MAKER' -- the text duplicate causes a double logo.
    """
    if not doc.tables:
        return 0

    table = doc.tables[0]  # Cover page is always table 0
    removed = 0

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                # Check each run for "MAKER" text that isn't part of a longer string
                for run in p.runs:
                    if run.text.strip() == "MAKER":
                        # Verify this is the standalone "MAKER" text, not "MAKER WEB STUDIOS"
                        run.text = ""
                        removed += 1

    return removed


def rebuild_footer(doc, logo_path, qr_path):
    """
    Rebuild the footer across all sections:
    Layout: [LOGO]  www.makerwebstudios.com | CONFIDENTIAL — PREPARED FOR [CLIENT NAME]  [QR CODE]
    Using a 3-column table for alignment.
    """
    for section in doc.sections:
        footer = section.footer
        if not footer:
            continue

        # Get existing footer text to extract client name placeholder
        existing_text = ""
        for p in footer.paragraphs:
            existing_text += p.text

        # Clear existing footer content
        for p in footer.paragraphs:
            p.clear()

        # Remove existing paragraphs except the first one
        while len(footer.paragraphs) > 1:
            p_elem = footer.paragraphs[-1]._element
            p_elem.getparent().remove(p_elem)

        # Clear the remaining paragraph
        first_para = footer.paragraphs[0]
        first_para.clear()

        # Create a table in the footer for layout: logo | text | qr
        tbl = doc.add_table(rows=1, cols=3)

        # Move the table into the footer
        tbl_elem = tbl._element
        tbl_elem.getparent().remove(tbl_elem)
        footer._element.append(tbl_elem)

        # Remove borders from the table
        tbl_props = tbl_elem.find(qn('w:tblPr'))
        if tbl_props is None:
            tbl_props = etree.SubElement(tbl_elem, qn('w:tblPr'))

        # Set table width to 100%
        tbl_width = tbl_props.find(qn('w:tblW'))
        if tbl_width is None:
            tbl_width = etree.SubElement(tbl_props, qn('w:tblW'))
        tbl_width.set(qn('w:w'), '5000')
        tbl_width.set(qn('w:type'), 'pct')

        # Remove all borders
        borders = etree.SubElement(tbl_props, qn('w:tblBorders'))
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = etree.SubElement(borders, qn(f'w:{border_name}'))
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')

        # Column 0: Logo (left-aligned)
        cell0 = tbl.cell(0, 0)
        p0 = cell0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if os.path.exists(logo_path):
            run0 = p0.add_run()
            run0.add_picture(logo_path, height=Inches(0.4))

        # Column 1: Text (left-aligned, sits next to logo)
        cell1 = tbl.cell(0, 1)
        p1 = cell1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run1 = p1.add_run("www.makerwebstudios.com  |  CONFIDENTIAL \u2014 PREPARED EXCLUSIVELY FOR [CLIENT NAME]")
        run1.font.size = Pt(7)
        run1.font.name = "Calibri"
        # Set vertical alignment to center
        tc_pr = cell1._element.find(qn('w:tcPr'))
        if tc_pr is None:
            tc_pr = etree.SubElement(cell1._element, qn('w:tcPr'))
        v_align = etree.SubElement(tc_pr, qn('w:vAlign'))
        v_align.set(qn('w:val'), 'center')

        # Column 2: QR Code (right-aligned)
        cell2 = tbl.cell(0, 2)
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if os.path.exists(qr_path):
            run2 = p2.add_run()
            run2.add_picture(qr_path, height=Inches(0.4))

        # Set column widths: logo narrow, text wide, QR narrow
        for row in tbl.rows:
            for ci, cell in enumerate(row.cells):
                tc_pr = cell._element.find(qn('w:tcPr'))
                if tc_pr is None:
                    tc_pr = etree.SubElement(cell._element, qn('w:tcPr'))
                tc_w = tc_pr.find(qn('w:tcW'))
                if tc_w is None:
                    tc_w = etree.SubElement(tc_pr, qn('w:tcW'))
                if ci == 0:  # Logo
                    tc_w.set(qn('w:w'), '1000')
                    tc_w.set(qn('w:type'), 'pct')
                elif ci == 1:  # Text
                    tc_w.set(qn('w:w'), '3000')
                    tc_w.set(qn('w:type'), 'pct')
                else:  # QR
                    tc_w.set(qn('w:w'), '1000')
                    tc_w.set(qn('w:type'), 'pct')


# Process all templates
print("Fixing covers and footers across all templates...\n")

for fname in sorted(os.listdir(TEMPLATE_DIR)):
    if not fname.endswith('.docx'):
        continue

    fpath = os.path.join(TEMPLATE_DIR, fname)
    doc = Document(fpath)

    # Fix 1: Remove double logo
    removed = remove_maker_text_from_cover(doc)

    # Fix 2: Rebuild footer
    rebuild_footer(doc, LOGO_PATH, QR_PATH)

    doc.save(fpath)
    print(f"  OK: {fname} (removed {removed} 'MAKER' text duplicates, footer rebuilt)")

print("\nDone. All templates updated.")
