#!/usr/bin/env python3
"""
MWS Audit Auto-Fill v3
=======================
Usage:
    python3 mws_audit_fill.py <client.json> [md_dir] [output_dir]

Two-layer system:
  Layer 1 -- JSON identity fields (client name, scores, market data)
  Layer 2 -- .md content files from Claude Code (audit findings, analysis)

Requirements:
    pip install python-docx
"""

import json
import sys
import os
import re
from docx import Document


TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")

TEMPLATE_FILES = [
    "MKT-AU01-MWS-2026-04.06-Digital Presence Audit TEMPLATE.docx",
    "MKT-AU02-MWS-2026-04.06-Blind Spot Audit TEMPLATE.docx",
    "MKT-PB02-MWS-2026-04.06-Revenue Growth Playbook TEMPLATE.docx",
    "SAL-PB01-MWS-2026-04.06-Sales Playbook TEMPLATE.docx",
    "MKT-MF01-MWS-2026-04.06-Market Analysis 5-Year Forecast TEMPLATE.docx",
]

# ── LAYER 1: JSON -> template placeholder mapping ────────────────────────────

PLACEHOLDER_MAP = {
    "[Client Name]": "CLIENT_NAME", "[CLIENT NAME]": "CLIENT_NAME",
    "[Client name]": "CLIENT_NAME", "[Full Client Name]": "CLIENT_NAME",
    "[Full Company Name]": "CLIENT_NAME", "[Client]": "CLIENT_NAME",
    "[client]": "CLIENT_NAME", "[CLIENT]": "CLIENT_NAME",
    "[Client Website]": "CLIENT_WEBSITE", "[client website]": "CLIENT_WEBSITE",
    "[Month Year]": "MONTH_YEAR", "[Industry]": "INDUSTRY",
    "[Client Location]": "LOCATION", "[Market Name]": "MARKET_NAME",
    "[Year]": "CURRENT_YEAR", "[X.X/10]": "SCORE_OVERALL",
    "[CMS]": "CMS_PLATFORM", "[Avatar]": "AVATAR", "[avatar]": "AVATAR",
    "[Contact Name]": "CONTACT_NAME", "[Offer Name]": "OFFER_NAME",
    "[Bonus Name]": "BONUS_NAME", "[Primary Deliverable]": "PRIMARY_DELIVERABLE",
    "[Outcome]": "OUTCOME", "[outcome]": "OUTCOME",
}

# ── LAYER 2: .md heading -> template placeholder mapping ─────────────────────
# The .md files use simple ## Headings. This map translates them to the exact
# placeholder strings in the DOCX templates. Write once, works for every client.

MD_SECTION_MAP = {

    # ── BLIND SPOT AUDIT ─────────────────────────────────────────────────────
    "Opening Strategic Framing":
        "[Opening strategic framing \u2014 why the communication gap is the core problem, not the product or service.]",
    "Closing Strategic Framing":
        "[Closing quote \u2014 what becomes possible when the operational excellence matches the digital presence.]",
    "Blind Spot 1 Title": "[Blind Spot #1 Title]",
    "Blind Spot 2 Title": "[Blind Spot #2 Title]",
    "Blind Spot 3 Title": "[Blind Spot #3 Title]",
    "Blind Spot 4 Title": "[Blind Spot #4 Title]",
    "Blind Spot 5 Title": "[Blind Spot #5 Title]",
    "Blind Spot 1": "[Describe the blind spot \u2014 what specific credential or advantage is being hidden or communicated generically. What is it costing in lost leads. Why does it persist. Be specific about the competitor winning because of this gap.]",
    "Blind Spot 2": "[Describe what buyer questions are going unanswered \u2014 the exact research queries, the decision-stage content missing, and which competitor is capturing those buyers instead.]",
    "Blind Spot 3": "[Describe which digital channels the buyer uses before they ever call \u2014 and which ones the client is completely absent from. Be specific about directories, platforms, and search behavior.]",
    "Blind Spot 4": "[Describe the friction points in the buyer experience \u2014 phone number placement, CTA language, RFQ form absence, live chat gap. Describe what the current experience communicates vs. what it should.]",
    "Blind Spot 5": "[Describe how the company\u2019s single most powerful differentiator is buried in body text instead of leading the page. What does a 30-second visitor see vs. what they should see?]",
    "Fix 1-1": "[Fix 1-1]", "Fix 1-2": "[Fix 1-2]", "Fix 1-3": "[Fix 1-3]", "Fix 1-4": "[Fix 1-4]",
    "Fix 2-1": "[Fix 2-1]", "Fix 2-2": "[Fix 2-2]", "Fix 2-3": "[Fix 2-3]", "Fix 2-4": "[Fix 2-4]",
    "Fix 3-1": "[Fix 3-1]", "Fix 3-2": "[Fix 3-2]", "Fix 3-3": "[Fix 3-3]", "Fix 3-4": "[Fix 3-4]",
    "Fix 4-1": "[Fix 4-1]", "Fix 4-2": "[Fix 4-2]", "Fix 4-3": "[Fix 4-3]", "Fix 4-4": "[Fix 4-4]",
    "Fix 5-1": "[Fix 5-1]", "Fix 5-2": "[Fix 5-2]", "Fix 5-3": "[Fix 5-3]", "Fix 5-4": "[Fix 5-4]",

    # ── DIGITAL PRESENCE AUDIT ───────────────────────────────────────────────
    "The Company": "[Describe the company \u2014 what they do, who they serve, location, years in operation.]",
    "The Credentials": "[Key certifications, memberships, facility specs, notable clients.]",
    "The Platform": "[CMS, CDN, analytics status, last meaningful update.]",
    "The Market Reality": "[Key competitors and the client\u2019s underutilized advantages.]",
    "Strategic Finding": "[Key strategic finding \u2014 one sentence capturing the defining opportunity.]",
    "Executive Summary": "[Overall assessment]",

    # DPA audit sections -- 7 categories x (description + 2 pass + 1 warn + 3 fail)
    "Design UX Description": "[Design UX Description]",
    "Design UX Pass 1 Title": "[Design UX Pass 1 Title]", "Design UX Pass 1": "[Design UX Pass 1]",
    "Design UX Pass 2 Title": "[Design UX Pass 2 Title]", "Design UX Pass 2": "[Design UX Pass 2]",
    "Design UX Warn 1 Title": "[Design UX Warn 1 Title]", "Design UX Warn 1": "[Design UX Warn 1]",
    "Design UX Fail 1 Title": "[Design UX Fail 1 Title]", "Design UX Fail 1": "[Design UX Fail 1]",
    "Design UX Fail 2 Title": "[Design UX Fail 2 Title]", "Design UX Fail 2": "[Design UX Fail 2]",
    "Design UX Fail 3 Title": "[Design UX Fail 3 Title]", "Design UX Fail 3": "[Design UX Fail 3]",

    "OnPage SEO Description": "[OnPage SEO Description]",
    "OnPage SEO Pass 1 Title": "[OnPage SEO Pass 1 Title]", "OnPage SEO Pass 1": "[OnPage SEO Pass 1]",
    "OnPage SEO Pass 2 Title": "[OnPage SEO Pass 2 Title]", "OnPage SEO Pass 2": "[OnPage SEO Pass 2]",
    "OnPage SEO Warn 1 Title": "[OnPage SEO Warn 1 Title]", "OnPage SEO Warn 1": "[OnPage SEO Warn 1]",
    "OnPage SEO Fail 1 Title": "[OnPage SEO Fail 1 Title]", "OnPage SEO Fail 1": "[OnPage SEO Fail 1]",
    "OnPage SEO Fail 2 Title": "[OnPage SEO Fail 2 Title]", "OnPage SEO Fail 2": "[OnPage SEO Fail 2]",
    "OnPage SEO Fail 3 Title": "[OnPage SEO Fail 3 Title]", "OnPage SEO Fail 3": "[OnPage SEO Fail 3]",

    "Technical SEO Description": "[Technical SEO Description]",
    "Technical SEO Pass 1 Title": "[Technical SEO Pass 1 Title]", "Technical SEO Pass 1": "[Technical SEO Pass 1]",
    "Technical SEO Pass 2 Title": "[Technical SEO Pass 2 Title]", "Technical SEO Pass 2": "[Technical SEO Pass 2]",
    "Technical SEO Warn 1 Title": "[Technical SEO Warn 1 Title]", "Technical SEO Warn 1": "[Technical SEO Warn 1]",
    "Technical SEO Fail 1 Title": "[Technical SEO Fail 1 Title]", "Technical SEO Fail 1": "[Technical SEO Fail 1]",
    "Technical SEO Fail 2 Title": "[Technical SEO Fail 2 Title]", "Technical SEO Fail 2": "[Technical SEO Fail 2]",
    "Technical SEO Fail 3 Title": "[Technical SEO Fail 3 Title]", "Technical SEO Fail 3": "[Technical SEO Fail 3]",

    "Content Description": "[Content Description]",
    "Content Pass 1 Title": "[Content Pass 1 Title]", "Content Pass 1": "[Content Pass 1]",
    "Content Pass 2 Title": "[Content Pass 2 Title]", "Content Pass 2": "[Content Pass 2]",
    "Content Warn 1 Title": "[Content Warn 1 Title]", "Content Warn 1": "[Content Warn 1]",
    "Content Fail 1 Title": "[Content Fail 1 Title]", "Content Fail 1": "[Content Fail 1]",
    "Content Fail 2 Title": "[Content Fail 2 Title]", "Content Fail 2": "[Content Fail 2]",
    "Content Fail 3 Title": "[Content Fail 3 Title]", "Content Fail 3": "[Content Fail 3]",

    "Conversion Description": "[Conversion Description]",
    "Conversion Pass 1 Title": "[Conversion Pass 1 Title]", "Conversion Pass 1": "[Conversion Pass 1]",
    "Conversion Pass 2 Title": "[Conversion Pass 2 Title]", "Conversion Pass 2": "[Conversion Pass 2]",
    "Conversion Warn 1 Title": "[Conversion Warn 1 Title]", "Conversion Warn 1": "[Conversion Warn 1]",
    "Conversion Fail 1 Title": "[Conversion Fail 1 Title]", "Conversion Fail 1": "[Conversion Fail 1]",
    "Conversion Fail 2 Title": "[Conversion Fail 2 Title]", "Conversion Fail 2": "[Conversion Fail 2]",
    "Conversion Fail 3 Title": "[Conversion Fail 3 Title]", "Conversion Fail 3": "[Conversion Fail 3]",

    "Trust Description": "[Trust Description]",
    "Trust Pass 1 Title": "[Trust Pass 1 Title]", "Trust Pass 1": "[Trust Pass 1]",
    "Trust Pass 2 Title": "[Trust Pass 2 Title]", "Trust Pass 2": "[Trust Pass 2]",
    "Trust Warn 1 Title": "[Trust Warn 1 Title]", "Trust Warn 1": "[Trust Warn 1]",
    "Trust Fail 1 Title": "[Trust Fail 1 Title]", "Trust Fail 1": "[Trust Fail 1]",
    "Trust Fail 2 Title": "[Trust Fail 2 Title]", "Trust Fail 2": "[Trust Fail 2]",
    "Trust Fail 3 Title": "[Trust Fail 3 Title]", "Trust Fail 3": "[Trust Fail 3]",

    "Digital Presence Description": "[Digital Presence Description]",
    "Digital Presence Pass 1 Title": "[Digital Presence Pass 1 Title]", "Digital Presence Pass 1": "[Digital Presence Pass 1]",
    "Digital Presence Pass 2 Title": "[Digital Presence Pass 2 Title]", "Digital Presence Pass 2": "[Digital Presence Pass 2]",
    "Digital Presence Warn 1 Title": "[Digital Presence Warn 1 Title]", "Digital Presence Warn 1": "[Digital Presence Warn 1]",
    "Digital Presence Fail 1 Title": "[Digital Presence Fail 1 Title]", "Digital Presence Fail 1": "[Digital Presence Fail 1]",
    "Digital Presence Fail 2 Title": "[Digital Presence Fail 2 Title]", "Digital Presence Fail 2": "[Digital Presence Fail 2]",
    "Digital Presence Fail 3 Title": "[Digital Presence Fail 3 Title]", "Digital Presence Fail 3": "[Digital Presence Fail 3]",

    # DPA strengths
    "Strength 1 Title": "[Strength 1 Title]", "Strength 1": "[Strength 1]",
    "Strength 2 Title": "[Strength 2 Title]", "Strength 2": "[Strength 2]",
    "Strength 3 Title": "[Strength 3 Title]", "Strength 3": "[Strength 3]",
    "Strength 4 Title": "[Strength 4 Title]", "Strength 4": "[Strength 4]",
    "Strength 5 Title": "[Strength 5 Title]", "Strength 5": "[Strength 5]",
    "Strength 6 Title": "[Strength 6 Title]", "Strength 6": "[Strength 6]",

    # DPA score notes
    "Design UX Score Notes": "[Design UX Score Notes]",
    "OnPage SEO Score Notes": "[OnPage SEO Score Notes]",
    "Technical SEO Score Notes": "[Technical SEO Score Notes]",
    "Content Score Notes": "[Content Score Notes]",
    "Conversion Score Notes": "[Conversion Score Notes]",
    "Trust Score Notes": "[Trust Score Notes]",
    "Digital Presence Score Notes": "[Digital Presence Score Notes]",

    # DPA priority fixes
    "Critical Fix 1": "[Critical fix #1 \u2014 what it is and exactly what to do.]",
    "Critical Fix 2": "[Critical fix #2]",
    "Critical Fix 3": "[Critical fix #3]",
    "High Fix 1": "[High priority fix #4]",
    "High Fix 2": "[High priority fix #5]",
    "High Fix 3": "[High priority fix #6]",
    "High Fix 4": "[High priority fix #7]",
    "Medium Fix 1": "[Medium priority fix #8]",
    "Medium Fix 2": "[Medium priority fix #9]",
    "Medium Fix 3": "[Medium priority fix #10]",

    # DPA roadmap
    "Foundation Task 1": "[Foundation task 1]", "Foundation Task 2": "[Foundation task 2]",
    "Foundation Task 3": "[Foundation task 3]", "Foundation Task 4": "[Foundation task 4]",
    "Content Task 1": "[Content task 1]", "Content Task 2": "[Content task 2]",
    "Content Task 3": "[Content task 3]", "Content Task 4": "[Content task 4]",
    "Growth Task 1": "[Growth task 1]", "Growth Task 2": "[Growth task 2]",
    "Growth Task 3": "[Growth task 3]", "Growth Task 4": "[Growth task 4]",

    # DPA competitors
    "Competitor 1 Strengths": "[Competitor 1 Strengths]",
    "Competitor 2 Strengths": "[Competitor 2 Strengths]",
    "Competitor 3 Strengths": "[Competitor 3 Strengths]",

    # DPA closing
    "Closing Strategic Quote": "[Closing strategic quote \u2014 one or two sentences that capture the opportunity with urgency and precision.]",
    "Key Competitive Finding": "[Key competitive context finding \u2014 why the market advantage is being lost digitally.]",

    # ── REVENUE GROWTH PLAYBOOK (ex-Grand Slam) ──────────────────────────────
    "Current Offer Problem": "[Describe the current offer language \u2014 the generic tagline, the audit score, the gap between world-class credentials and a low-scoring website. Be specific.]",
    "Avatar Detail": "[Job title, company type, revenue size, geography, situation \u2014 be specific.]",
    "Primary Markets": "[Primary markets \u2014 where the highest-value clients are located or searching from.]",
    "Dream Outcome 1": "[Dream outcome #1 \u2014 specific result, not a feature]",
    "Dream Outcome 2": "[Dream outcome #2]", "Dream Outcome 3": "[Dream outcome #3]",
    "Dream Outcome 4": "[Dream outcome #4]", "Dream Outcome 5": "[Dream outcome #5]",
    "Offer Rename": "[Generic service name \u2014 what the client currently calls it]",
    "Value Stack Core": "[Core deliverable 1]",
    "Guarantee Layer 1": "[Unconditional guarantee statement \u2014 remove all buyer risk in the first 30 days.]",
    "Guarantee Layer 2": "[Conditional guarantee \u2014 specific measurable outcome, or a fee is waived.]",
    "Guarantee Rationale": "[Buyer\u2019s #1 fear and how each guarantee layer directly addresses it.]",
    "Scarcity": "[What is genuinely limited? Real operational constraints \u2014 floor space, production lines, slots per quarter, territory. Never manufacture scarcity.]",
    "Urgency": "[What market force is happening NOW that makes waiting costly? Tariffs, competitor capacity lock-up, seasonal windows, regulatory changes.]",
    "What They Are Sitting On": "[Asset/Credential Title]",
    "Competitive Moat": "[Describe the fundamental structural difference between this client and every competitor. What do competitors do? What IS this client? That category distinction is the moat.]",
    "The Pitch": "[Final summary: what the client is vs. how they are presenting. The  framework communicates what the business already is. Package it, guarantee it, and name it.]",
    "One Line Differentiator": "[The one-line differentiator \u2014 the sentence that could not be said by any competitor in the market.]",
    "Positioning Statement": "[One sentence positioning statement \u2014 what the client actually IS, not what their tagline says.]",
    "Market Indicators": "[Describe why the client\u2019s market passes this test \u2014 use specific data.]",
    # Website changes table
    "Hero Current": "[Hero Current]", "Hero Recommended": "[Hero Recommended]",
    "Sub-hero Current": "[Sub-hero Current]", "Sub-hero Recommended": "[Sub-hero Recommended]",
    "CTA Current": "[CTA Current]", "CTA Recommended": "[CTA Recommended]",
    "Trust Badges Current": "[Trust Badges Current]", "Trust Badges Recommended": "[Trust Badges Recommended]",
    "Services Current": "[Services Current]", "Services Recommended": "[Services Recommended]",
    "Case Studies Current": "[Case Studies Current]", "Case Studies Recommended": "[Case Studies Recommended]",
    "Guarantee Current": "[Guarantee Current]", "Guarantee Recommended": "[Guarantee Recommended]",
    "Lead Capture Current": "[Lead Capture Current]", "Lead Capture Recommended": "[Lead Capture Recommended]",
    "Phone Current": "[Phone Current]", "Phone Recommended": "[Phone Recommended]",
    "Bottom CTA Current": "[Bottom CTA Current]", "Bottom CTA Recommended": "[Bottom CTA Recommended]",
    "Complete Offer Pitch": "[Complete offer pitch: (1) The mistake buyers make without this. (2) What client has already built and solved. (3) Core credentials. (4) The guarantee. (5) The scarcity. One paragraph.]",

    # ── GROWTH PLAYBOOK ──────────────────────────────────────────────────────
    "Growth Buyer Profile": "[Growth Buyer Profile]",
    "Growth Pain Description": "[Growth Pain Description]",
    "Growth Buyer Description": "[Buyer description]",
    "Growth Problem Description": "[Problem description]",
    # Services table
    "Service 1 Buyer": "[Service 1 Buyer]", "Service 1 Demand": "[Service 1 Demand]",
    "Service 2 Buyer": "[Service 2 Buyer]", "Service 2 Demand": "[Service 2 Demand]",
    "Service 3 Buyer": "[Service 3 Buyer]", "Service 3 Demand": "[Service 3 Demand]",
    "Service 4 Buyer": "[Service 4 Buyer]", "Service 4 Demand": "[Service 4 Demand]",
    # ROI ranking
    "ROI Client 1": "[ROI Client 1]", "ROI Reasoning 1": "[ROI Reasoning 1]",
    "ROI Client 2": "[ROI Client 2]", "ROI Reasoning 2": "[ROI Reasoning 2]",
    "ROI Client 3": "[ROI Client 3]", "ROI Reasoning 3": "[ROI Reasoning 3]",
    "ROI Client 4": "[ROI Client 4]", "ROI Reasoning 4": "[ROI Reasoning 4]",
    "ROI Client 5": "[ROI Client 5]", "ROI Reasoning 5": "[ROI Reasoning 5]",
    # Skills
    "Skill 1 Title": "[Skill 1 Title]", "Skill 1 Score": "[Skill 1 Score]",
    "Skill 2 Title": "[Skill 2 Title]", "Skill 2 Score": "[Skill 2 Score]",
    "Skill 3 Title": "[Skill 3 Title]",
    "Skill 4 Title": "[Skill 4 Title]", "Skill 4 Score": "[Skill 4 Score]",
    "Skill 5 Title": "[Skill 5 Title]",
    # Actions
    "Action 1 Title": "[Action 1 Title]", "Action 1 Description": "[Action 1 Description]",
    "Action 2 Title": "[Action 2 Title]", "Action 2 Description": "[Action 2 Description]",
    "Action 3 Title": "[Action 3 Title]", "Action 3 Description": "[Action 3 Description]",
    "Action 4 Title": "[Action 4 Title]", "Action 4 Description": "[Action 4 Description]",
    "Action 5 Title": "[Action 5 Title]", "Action 5 Description": "[Action 5 Description]",

    # ── MARKET ANALYSIS ──────────────────────────────────────────────────────
    "Regional Factor 1": "[Regional Factor 1 \u2014 e.g. Agricultural Export Corridor]",
    "Regional Factor 1 Detail": "[Regional Factor 1 Detail]",
    "Regional Factor 2": "[Regional Factor 2 \u2014 e.g. Industrial / Manufacturing Base]",
    "Regional Factor 2 Detail": "[Regional Factor 2 Detail]",
    "Regional Factor 3": "[Regional Factor 3 \u2014 e.g. Construction & Commercial Growth]",
    "Regional Factor 3 Detail": "[Regional Factor 3 Detail]",
    "Driver 1 Title": "[Driver 1 Title]",
    "Driver 1 Detail": "[Describe this structural trend: CAGR, market size, why it is growing, and how it creates demand for the client\u2019s specific product or service. End with a specific insight about the client\u2019s positioning advantage. 2\u20133 sentences.]",
    "Driver 2 Title": "[Driver 2 Title]",
    "Driver 3 Title": "[Driver 3 Title]",
    "Driver 4 Title": "[Driver 4 Title]",
    "Driver 5 Title": "[Driver 5 Title]",
    "Location Advantage": "[Location Advantage]",
    "Risk 1": "[Risk 1]", "Risk 1 Mitigation": "[Risk 1 Mitigation]",
    "Risk 2": "[Risk 2]", "Risk 2 Mitigation": "[Risk 2 Mitigation]",
    "Risk 3": "[Risk 3]", "Risk 3 Mitigation": "[Risk 3 Mitigation]",
    "Risk 4": "[Risk 4]", "Risk 4 Mitigation": "[Risk 4 Mitigation]",
    "Risk 5": "[Risk 5]", "Risk 5 Mitigation": "[Risk 5 Mitigation]",
    "Forecast 2026 Assumption": "[Forecast 2026 Assumption]",
    "Forecast 2027 Assumption": "[Forecast 2027 Assumption]",
    "Forecast 2028 Assumption": "[Forecast 2028 Assumption]",
    "Forecast 2029 Assumption": "[Forecast 2029 Assumption]",
    "Forecast 2030 Assumption": "[Forecast 2030 Assumption]",
    "Timeline 2026 Milestone": "[Timeline 2026 Milestone]",
    "Timeline 2027 Milestone": "[Timeline 2027 Milestone]",
    "Timeline 2028 Milestone": "[Timeline 2028 Milestone]",
    "Timeline 2029 Milestone": "[Timeline 2029 Milestone]",
    "Timeline 2030 Milestone": "[Timeline 2030 Milestone]",
    "Strategic Horizon": "[Write 1\u20132 paragraphs on what market leadership looks like for this client in 5 years \u2014 exit valuation, client roster, brand recognition, second facility. What must be true about their digital presence for that to occur?]",

    # ── REVENUE ACCELERATION PLAYBOOK ────────────────────────────────────────
    "Commodity Problem": "[Describe how the client currently presents their offer \u2014 and why it sounds like every other competitor. Reference their homepage tagline and audit score.]",
    "What They Actually Are": "[Cut through the generic tagline. List the real credentials: years in operation, facility size, certifications, geography. These are the raw materials of the Grand Slam Offer.]",
    "Headline Statement": "[One sentence that could be the homepage headline \u2014 what the client IS, not what they claim to be.]",
    "Core Deliverable": "[Describe the core offer \u2014 what is delivered, what is included, anchored annual value range.]",
    "Bonus 1": "[Bonus 1 Description]",
    "Bonus 2": "[Bonus 2 Description]",
    "Bonus 3": "[Bonus 3 Description]",
    "RA Guarantee Layer 1": "[Write the unconditional guarantee. What can the client promise in the first 30 days that removes all buyer doubt? E.g., a facility visit with a full refund if it doesn\u2019t match what was promised.]",
    "RA Guarantee Layer 2": "[Write the conditional guarantee. Specific measurable outcome in 90 days \u2014 or a fee is waived. E.g., first production run ships in 90 days or first month\u2019s shelter fee is waived.]",
    "RA Guarantee Rationale": "[Describe the buyer\u2019s #1 fear about this purchase \u2014 and explain how each guarantee layer directly addresses it.]",
    "RA Scarcity": "[What is genuinely limited? Facility capacity, slots per quarter, geographic availability, production line capacity. Use real operational constraints \u2014 never manufacture scarcity.]",
    "RA Urgency": "[What market force is happening NOW that makes waiting costly? Tariff changes, competitor capacity lock-up, seasonal windows, regulatory deadlines, rising input costs.]",
    "RA Pitch": "[Complete offer pitch: (1) The mistake buyers make without this. (2) What client has already solved. (3) Core credentials. (4) The guarantee. (5) The scarcity. One paragraph \u2014 read it aloud, it should take 20\u201325 seconds.]",
    "Stat 2 Label": "[Stat 2 Label]",
    "Stat 3 Label": "[Stat 3 Label]",
    "Stat 4 Cert": "[Stat 4 Cert]",
    "Stat 4 Abbrev": "[Stat 4 Abbrev]",
}

# Strategic Opportunity maps to multiple templates -- add both
MD_SECTION_MAP["Strategic Opportunity"] = "[Write 2\u20133 paragraphs: what the client has earned through operational execution, why these blind spots are communication problems not product problems, and what becomes possible when they are closed.]"


# ── ENGINE ───────────────────────────────────────────────────────────────────

def load_client(json_path):
    with open(json_path, "r", encoding="utf-8") as f:
        return json.load(f)


def build_json_replacements(client):
    r = {}
    for ph, key in PLACEHOLDER_MAP.items():
        v = client.get(key, "")
        if v:
            r[ph] = str(v)
    for key, val in client.items():
        if key.startswith("_"):
            continue
        sv = str(val)
        if sv:
            r[f"[{key}]"] = sv
            r[f"[{key.lower()}]"] = sv
            r[f"[{key.title()}]"] = sv
    return r


def parse_md_sections(md_path):
    sections = {}
    heading = None
    lines = []
    with open(md_path, "r", encoding="utf-8") as f:
        for line in f:
            m = re.match(r'^##\s+(.+)$', line.strip())
            if m:
                if heading:
                    sections[heading] = "\n".join(lines).strip()
                heading = m.group(1).strip()
                lines = []
            elif heading:
                lines.append(line.rstrip())
    if heading:
        sections[heading] = "\n".join(lines).strip()
    return sections


def parse_skill_deliverables(all_sections):
    """
    Parse pipe-delimited deliverable lines from ## Skill N sections.
    Format: 'Deliverable N: Name | Included | Outcome'
    Creates mappings like [Skill 1 Deliverable 1], [Skill 1 Included 1], etc.
    """
    r = {}
    for skill_n in range(1, 6):
        key = f"Skill {skill_n}"
        if key not in all_sections:
            continue
        content = all_sections[key]
        for line in content.split("\n"):
            m = re.match(r'Deliverable\s+(\d+):\s*(.+)', line.strip())
            if m:
                dn = m.group(1)
                parts = [p.strip() for p in m.group(2).split("|")]
                if len(parts) >= 1:
                    r[f"[Skill {skill_n} Deliverable {dn}]"] = parts[0]
                if len(parts) >= 2:
                    r[f"[Skill {skill_n} Included {dn}]"] = parts[1]
                    r[f"[What's included]"] = parts[1]  # fallback for T4
                if len(parts) >= 3:
                    r[f"[Skill {skill_n} Outcome {dn}]"] = parts[2]
    return r


def build_md_replacements(all_sections):
    r = {}
    for md_heading, template_ph in MD_SECTION_MAP.items():
        if md_heading in all_sections and all_sections[md_heading]:
            r[template_ph] = all_sections[md_heading]
    # Parse skill deliverables from pipe-delimited format
    r.update(parse_skill_deliverables(all_sections))
    # Direct [heading] fallback
    for heading, content in all_sections.items():
        if content:
            r[f"[{heading}]"] = content
    return r


def replace_in_text(text, replacements):
    for ph, val in replacements.items():
        text = text.replace(ph, val)
    return text


def replace_in_paragraph(para, replacements):
    full = "".join(run.text for run in para.runs)
    new = replace_in_text(full, replacements)
    if new != full and para.runs:
        # Find the run with the most text -- this has the "dominant" formatting
        # (avoids splitting replacement text across runs with different colors)
        dominant_idx = max(range(len(para.runs)), key=lambda i: len(para.runs[i].text))
        # Put ALL replacement text in the dominant run, clear all others
        for i, run in enumerate(para.runs):
            if i == dominant_idx:
                run.text = new
            else:
                run.text = ""
    else:
        for run in para.runs:
            if run.text:
                run.text = replace_in_text(run.text, replacements)


def replace_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_in_paragraph(para, replacements)
            for nested in cell.tables:
                replace_in_table(nested, replacements)


def replace_in_document(doc, replacements):
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)
    for table in doc.tables:
        replace_in_table(table, replacements)
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                replace_in_paragraph(para, replacements)
            for table in section.header.tables:
                replace_in_table(table, replacements)
        if section.footer:
            for para in section.footer.paragraphs:
                replace_in_paragraph(para, replacements)
            for table in section.footer.tables:
                replace_in_table(table, replacements)


def strip_remaining_brackets(doc):
    bracket_re = re.compile(r'\[[^\]]{3,}\]')
    def clean(para):
        full = "".join(r.text for r in para.runs)
        new = bracket_re.sub('', full)
        new = re.sub(r'  +', ' ', new).strip()
        if new != full and para.runs:
            remaining = new
            for i, run in enumerate(para.runs):
                if i < len(para.runs) - 1:
                    n = len(run.text)
                    run.text = remaining[:n]
                    remaining = remaining[n:]
                else:
                    run.text = remaining

    for para in doc.paragraphs:
        clean(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    clean(para)
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            for para in ncell.paragraphs:
                                clean(para)
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                clean(para)
        if section.footer:
            for para in section.footer.paragraphs:
                clean(para)


def sanitize_filename(name):
    return re.sub(r'[^\w\s-]', '', name).strip().replace(' ', '_')


def process_client(json_path, md_dir, output_dir):
    print(f"\n{'='*60}")
    print(f"MWS Audit Auto-Fill v3")
    print(f"{'='*60}")

    client = load_client(json_path)
    client_name = client.get("CLIENT_NAME", "Unknown_Client")
    client_slug = sanitize_filename(client_name)
    replacements = build_json_replacements(client)
    print(f"Client: {client_name}")
    print(f"JSON replacements: {len(replacements)}")

    all_md = {}
    if md_dir and os.path.isdir(md_dir):
        for md_file in sorted(os.listdir(md_dir)):
            if md_file.endswith(".md"):
                sections = parse_md_sections(os.path.join(md_dir, md_file))
                all_md.update(sections)
                print(f"  .md: {md_file} ({len(sections)} sections)")
        md_r = build_md_replacements(all_md)
        replacements.update(md_r)
        print(f"Total replacements: {len(replacements)}")
    else:
        print("No .md directory -- JSON-only mode")

    print(f"{'~'*60}")
    os.makedirs(output_dir, exist_ok=True)
    generated = []

    for tf in TEMPLATE_FILES:
        tp = os.path.join(TEMPLATE_DIR, tf)
        if not os.path.exists(tp):
            print(f"  SKIP: {tf}")
            continue
        dt = re.sub(r'^[A-Z]{3}-[A-Z]{2}\d{2}-MWS-[\d.]+-', '', tf).replace(" TEMPLATE", "").replace(".docx", "")
        of = f"{client_slug}_{dt}.docx"
        op = os.path.join(output_dir, of)
        try:
            doc = Document(tp)
            replace_in_document(doc, replacements)
            strip_remaining_brackets(doc)
            doc.save(op)
            generated.append(of)
            print(f"  OK  {of}")
        except Exception as e:
            print(f"  ERR {tf}: {e}")

    print(f"\n{'~'*60}")
    print(f"Generated {len(generated)} docs for {client_name}")
    print(f"Output: {output_dir}")
    print(f"{'='*60}\n")
    return generated


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 mws_audit_fill.py <client.json> [md_dir] [output_dir]")
        sys.exit(1)
    json_path = sys.argv[1]
    md_dir = sys.argv[2] if len(sys.argv) > 2 else None
    output_dir = sys.argv[3] if len(sys.argv) > 3 else "output"
    process_client(json_path, md_dir, output_dir)
