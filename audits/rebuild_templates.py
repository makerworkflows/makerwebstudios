#!/usr/bin/env python3
"""
Rebuild all 6 DOCX templates with unique placeholders per cell.
Run once to update templates, then delete this script.
"""
from docx import Document
import os, re

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")


def replace_para_text(para, old, new):
    """Replace text in a paragraph preserving first run's formatting."""
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False
    new_text = full.replace(old, new)
    if para.runs:
        remaining = new_text
        for i, run in enumerate(para.runs):
            if i < len(para.runs) - 1:
                orig_len = len(run.text)
                run.text = remaining[:orig_len]
                remaining = remaining[orig_len:]
            else:
                run.text = remaining
    return True


def replace_cell_text(cell, old, new):
    """Replace text in all paragraphs of a table cell."""
    for p in cell.paragraphs:
        replace_para_text(p, old, new)


def full_cell_text(cell):
    return "\n".join(p.text for p in cell.paragraphs)


# ═══════════════════════════════════════════════════════════════════════════════
# 1. DIGITAL PRESENCE AUDIT
# ═══════════════════════════════════════════════════════════════════════════════
print("Rebuilding: Digital Presence Audit...")
doc = Document(os.path.join(TEMPLATE_DIR, "MKT-AU01-MWS-2026-04.06-Digital Presence Audit TEMPLATE.docx"))

CATEGORIES = ["Design UX", "OnPage SEO", "Technical SEO", "Content", "Conversion", "Trust", "Digital Presence"]
cat_idx = -1
pass_n = warn_n = fail_n = 0

for i, p in enumerate(doc.paragraphs):
    full = "".join(r.text for r in p.runs)

    # Detect category section headers
    for ci, name in enumerate(["Design & UX Audit", "On-Page SEO Audit", "Technical SEO Audit",
                                "Content Quality Audit", "Conversion Audit", "Trust & Credibility Audit",
                                "Digital Presence Audit"]):
        if name in full:
            # For "Digital Presence Audit" make sure it's section 10 not the doc title
            if ci == 6 and i < 30:
                continue
            cat_idx = ci
            pass_n = warn_n = fail_n = 0
            break

    if cat_idx < 0:
        continue
    cat = CATEGORIES[cat_idx]

    if "[Description of what this category covers for the client.]" in full:
        replace_para_text(p, "[Description of what this category covers for the client.]", f"[{cat} Description]")

    if "[Pass item]" in full:
        pass_n += 1
        replace_para_text(p, "[Pass item]", f"[{cat} Pass {pass_n} Title]")
        full2 = "".join(r.text for r in p.runs)
        replace_para_text(p, "[What is working and why it matters to buyers or search engines.]", f"[{cat} Pass {pass_n}]")
        full3 = "".join(r.text for r in p.runs)
        replace_para_text(p, "[What is working.]", f"[{cat} Pass {pass_n}]")

    if "[Warning item]" in full:
        warn_n += 1
        replace_para_text(p, "[Warning item]", f"[{cat} Warn {warn_n} Title]")
        replace_para_text(p, "[What needs improvement but is not a critical failure.]", f"[{cat} Warn {warn_n}]")

    if "[Fail item]" in full:
        fail_n += 1
        replace_para_text(p, "[Fail item]", f"[{cat} Fail {fail_n} Title]")
        full2 = "".join(r.text for r in p.runs)
        if "[What is failing \u2014 include the revenue, trust, or SEO impact.]" in full2:
            replace_para_text(p, "[What is failing \u2014 include the revenue, trust, or SEO impact.]", f"[{cat} Fail {fail_n}]")
        full3 = "".join(r.text for r in p.runs)
        if "[What is failing.]" in full3:
            replace_para_text(p, "[What is failing.]", f"[{cat} Fail {fail_n}]")

# Strength tables -- find by numbered cells
strength_num = 0
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = full_cell_text(cell)
            if "[Strength Title]" in txt:
                strength_num += 1
                replace_cell_text(cell, "[Strength Title]", f"[Strength {strength_num} Title]")
                replace_cell_text(cell, "[Describe what they are doing right and why it matters to buyers or search engines.]",
                                 f"[Strength {strength_num}]")

# Score notes
notes_num = 0
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if cell.text.strip() == "[Notes]":
                notes_num += 1
                cat_name = CATEGORIES[notes_num - 1] if notes_num <= len(CATEGORIES) else f"Cat{notes_num}"
                for p in cell.paragraphs:
                    if p.text.strip() == "[Notes]" and p.runs:
                        p.runs[0].text = f"[{cat_name} Score Notes]"

# Competitor descriptions
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            replace_cell_text(cell, "[Description of their digital strengths vs. the client]", "[Competitor 1 Strengths]")
            replace_cell_text(cell, "[Description of their digital strengths]", "[Competitor 2 Strengths]")
            replace_cell_text(cell, "[Their real credentials vs. digital underperformance]", "[Competitor 3 Strengths]")

doc.save(os.path.join(TEMPLATE_DIR, "MKT-AU01-MWS-2026-04.06-Digital Presence Audit TEMPLATE.docx"))
print("  OK: Digital Presence Audit")


# ═══════════════════════════════════════════════════════════════════════════════
# 2. GRAND SLAM OFFER -> REVENUE GROWTH PLAYBOOK
# ═══════════════════════════════════════════════════════════════════════════════
print("Rebuilding: Grand Slam Offer -> Revenue Growth Playbook...")
doc = Document(os.path.join(TEMPLATE_DIR, "SAL-PP01-MWS-2026-04.06-Grand Slam Offer Framework TEMPLATE.docx"))

# Rename everywhere
for p in doc.paragraphs:
    replace_para_text(p, "GRAND SLAM OFFER FRAMEWORK", "REVENUE GROWTH PLAYBOOK")
    replace_para_text(p, "Grand Slam Offer Framework", "Revenue Growth Playbook")
    replace_para_text(p, "Grand Slam Offer", "Revenue Growth Playbook")
    replace_para_text(p, "grand slam offer", "revenue growth playbook")
    replace_para_text(p, "Based on Alex Hormozi\u2019s $100M Offers", "Strategic Offer Architecture for Maximum Revenue Impact")
    replace_para_text(p, "Based on Alex Hormozi's $100M Offers", "Strategic Offer Architecture for Maximum Revenue Impact")
    replace_para_text(p, "Alex Hormozi", "")
    replace_para_text(p, "$100M Offers", "")
    replace_para_text(p, "$100M", "")
    replace_para_text(p, "Hormozi", "")
    replace_para_text(p, "HORMOZI VERSION", "RECOMMENDED VERSION")
    replace_para_text(p, "Hormozi version", "Recommended version")

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            replace_cell_text(cell, "GRAND SLAM OFFER FRAMEWORK", "REVENUE GROWTH PLAYBOOK")
            replace_cell_text(cell, "Grand Slam Offer Framework", "Revenue Growth Playbook")
            replace_cell_text(cell, "Grand Slam Offer", "Revenue Growth Playbook")
            replace_cell_text(cell, "Based on Alex Hormozi\u2019s $100M Offers", "Strategic Offer Architecture for Maximum Revenue Impact")
            replace_cell_text(cell, "Based on Alex Hormozi's $100M Offers", "Strategic Offer Architecture for Maximum Revenue Impact")
            replace_cell_text(cell, "Alex Hormozi", "")
            replace_cell_text(cell, "$100M Offers", "")
            replace_cell_text(cell, "$100M", "")
            replace_cell_text(cell, "Hormozi", "")
            replace_cell_text(cell, "HORMOZI VERSION", "RECOMMENDED VERSION")
            replace_cell_text(cell, "Hormozi version", "Recommended version")

# Headers/footers
for section in doc.sections:
    if section.header:
        for p in section.header.paragraphs:
            replace_para_text(p, "GRAND SLAM OFFER FRAMEWORK", "REVENUE GROWTH PLAYBOOK")
    if section.footer:
        for p in section.footer.paragraphs:
            replace_para_text(p, "GRAND SLAM OFFER FRAMEWORK", "REVENUE GROWTH PLAYBOOK")

# Unique placeholders for repeating table cells
# Market indicators table (T2) - 4 rows with same placeholder
market_row = 0
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = full_cell_text(cell)
            if "[Describe why the client\u2019s market passes this test" in txt:
                market_row += 1
                labels = ["Massive Pain", "Purchasing Power", "Easy to Target", "Growing Market"]
                label = labels[market_row - 1] if market_row <= 4 else f"Indicator {market_row}"
                replace_cell_text(cell,
                    "[Describe why the client\u2019s market passes this test \u2014 use specific data.]",
                    f"[Market Indicator {label}]")

# Website changes table (T3) - 10 rows with same current/recommended
change_row = 0
WEBSITE_ELEMENTS = ["Hero", "Sub-hero", "CTA", "Trust Badges", "Services",
                     "Case Studies", "Guarantee", "Lead Capture", "Phone", "Bottom CTA"]
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        cells_text = [full_cell_text(c) for c in row.cells]
        if any("[What it currently says/does]" in t for t in cells_text):
            change_row += 1
            elem = WEBSITE_ELEMENTS[change_row - 1] if change_row <= 10 else f"Element {change_row}"
            for ci, cell in enumerate(row.cells):
                replace_cell_text(cell, "[What it currently says/does]", f"[{elem} Current]")
                replace_cell_text(cell, "[What it should say/do]", f"[{elem} Recommended]")

# Save with new name
new_path = os.path.join(TEMPLATE_DIR, "MKT-PB02-MWS-2026-04.06-Revenue Growth Playbook TEMPLATE.docx")
doc.save(new_path)
# Remove old file
old_path = os.path.join(TEMPLATE_DIR, "SAL-PP01-MWS-2026-04.06-Grand Slam Offer Framework TEMPLATE.docx")
if os.path.exists(old_path):
    os.remove(old_path)
print("  OK: Revenue Growth Playbook (renamed + cleaned)")


# ═══════════════════════════════════════════════════════════════════════════════
# 3. GROWTH PLAYBOOK - unique service rows, ROI rows, skills, outreach, actions
# ═══════════════════════════════════════════════════════════════════════════════
print("Rebuilding: Growth Playbook...")
doc = Document(os.path.join(TEMPLATE_DIR, "MKT-PB02-MWS-2026-04.06-Revenue Growth Playbook TEMPLATE.docx"))

# Services table (T1) - 4 rows
svc_row = 0
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        cells_text = [full_cell_text(c) for c in row.cells]
        if any("[Service]" in t for t in cells_text):
            svc_row += 1
            for ci, cell in enumerate(row.cells):
                replace_cell_text(cell, "[Service]", f"[Service {svc_row}]")
                replace_cell_text(cell, "[Buyer type]", f"[Service {svc_row} Buyer]")
                replace_cell_text(cell, "[HIGHEST / HIGH / MEDIUM]", f"[Service {svc_row} Demand]")

# ROI ranking table (T2) - 5 rows
roi_row = 0
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        cells_text = [full_cell_text(c) for c in row.cells]
        if any("[Client type / vertical]" in t for t in cells_text):
            roi_row += 1
            for ci, cell in enumerate(row.cells):
                replace_cell_text(cell, "[Client type / vertical]", f"[ROI Client {roi_row}]")
                replace_cell_text(cell, "[Why this client type yields the highest ROI for this business.]",
                                 f"[ROI Reasoning {roi_row}]")

# Skills tables - find [Capability Title] and [X/10]
skill_num = 0
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = full_cell_text(cell)
            if "[Capability Title]" in txt:
                skill_num += 1
                replace_cell_text(cell, "[Capability Title]", f"[Skill {skill_num} Title]")
                replace_cell_text(cell, "[X/10]", f"[Skill {skill_num} Score]")

# Deliverable tables - need to track which skill they belong to
deliv_skill = 0
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        cells_text = [full_cell_text(c) for c in row.cells]
        # Check if this table's first row has a skill reference or deliverable headers
        if ri == 0 and any("Deliverable" in t for t in cells_text):
            deliv_skill += 1
        if any("[Deliverable" in t for t in cells_text):
            dn = 0
            for ci, cell in enumerate(row.cells):
                txt = full_cell_text(cell)
                for d in ["[Deliverable 1]", "[Deliverable 2]", "[Deliverable 3]"]:
                    if d in txt:
                        dn = int(d.split()[1].rstrip("]"))
                        replace_cell_text(cell, d, f"[Skill {deliv_skill} Deliverable {dn}]")
                replace_cell_text(cell, "[What\u2019s included]", f"[Skill {deliv_skill} Included {dn}]")
                replace_cell_text(cell, "[Included]", f"[Skill {deliv_skill} Included {dn}]")
                replace_cell_text(cell, "[Outcome for client]", f"[Skill {deliv_skill} Outcome {dn}]")
                replace_cell_text(cell, "[Outcome]", f"[Skill {deliv_skill} Outcome {dn}]")

# Outreach emails - 6 sets in paragraphs
outreach_num = 0
for i, p in enumerate(doc.paragraphs):
    full = "".join(r.text for r in p.runs)
    if "[Prospect Type]" in full:
        outreach_num += 1
        replace_para_text(p, "[Prospect Type]", f"[Outreach {outreach_num} Prospect]")
        replace_para_text(p, "[Title]", f"[Outreach {outreach_num} Title]")
    if "[Email subject line" in full:
        replace_para_text(p, "[Email subject line \u2014 lead with their pain, not your service name]",
                         f"[Outreach {outreach_num} Subject]")
    if "[Sign-off]" in full:
        replace_para_text(p, "[Name]", f"[Outreach {outreach_num} Name]")
        replace_para_text(p, "[One sentence referencing their specific situation]", f"[Outreach {outreach_num} Line 1]")
        replace_para_text(p, "[Bridge from their pain to the credential that solves it]", f"[Outreach {outreach_num} Line 2]")
        replace_para_text(p, "[One specific proof point]", f"[Outreach {outreach_num} Line 3]")
        replace_para_text(p, "[Lowest-friction CTA]", f"[Outreach {outreach_num} CTA]")
        replace_para_text(p, "[Sign-off]", f"[Outreach {outreach_num} Signoff]")

# Action items - 5 in tables
action_num = 0
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = full_cell_text(cell)
            for a in [f"[Action {n} Title]" for n in range(1, 6)]:
                if a in txt:
                    n = int(a.split()[1])
                    replace_cell_text(cell, a, f"[Action {n} Title]")  # already unique
                    replace_cell_text(cell, "[Describe the specific action and why it creates the most leverage at this stage.]",
                                    f"[Action {n} Description]")

# Buyer/outcome descriptions
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            replace_cell_text(cell, "[Buyer description]", "[Growth Buyer Description]")
            replace_cell_text(cell, "[Problem description]", "[Growth Problem Description]")

# Paragraph-level buyer/pain descriptions
for p in doc.paragraphs:
    replace_para_text(p, "[Describe who the buyer is \u2014 job title, company type, revenue range, situation driving urgency.]",
                     "[Growth Buyer Profile]")
    replace_para_text(p, "[Describe the specific pain this capability eliminates. Why do buyers urgently need it? What happens without it?]",
                     "[Growth Pain Description]")

doc.save(os.path.join(TEMPLATE_DIR, "MKT-PB02-MWS-2026-04.06-Revenue Growth Playbook TEMPLATE.docx"))
print("  OK: Growth Playbook")


# ═══════════════════════════════════════════════════════════════════════════════
# 4. MARKET ANALYSIS - unique drivers, regional factors, risks, milestones
# ═══════════════════════════════════════════════════════════════════════════════
print("Rebuilding: Market Analysis...")
doc = Document(os.path.join(TEMPLATE_DIR, "MKT-MF01-MWS-2026-04.06-Market Analysis 5-Year Forecast TEMPLATE.docx"))

# Driver tables - 5 drivers with same description placeholder
driver_num = 0
DRIVER_PLACEHOLDERS = [
    "[Driver 1 \u2014 e.g. E-Commerce Fulfillment Growth]",
    "[Driver 2 \u2014 e.g. Nearshoring & Manufacturing Investment]",
    "[Driver 3 \u2014 e.g. Sustainability / Regulatory Mandates]",
    "[Driver 4 \u2014 e.g. Dominant End-User Segment]",
    "[Driver 5 \u2014 e.g. Premiumization / Niche Growth]",
]
DRIVER_DESC = "[Describe this structural trend: CAGR, market size, why it is growing, and how it creates demand for the client\u2019s specific product or service. End with a specific insight about the client\u2019s positioning advantage. 2\u20133 sentences.]"

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = full_cell_text(cell)
            for di, dp in enumerate(DRIVER_PLACEHOLDERS):
                if dp in txt:
                    replace_cell_text(cell, dp, f"[Driver {di+1} Title]")
            if DRIVER_DESC in txt:
                driver_num += 1
                replace_cell_text(cell, DRIVER_DESC, f"[Driver {driver_num} Detail]")

# Regional factors - already unique titles, but description repeats
REGIONAL_DESC = "[Describe this regional factor: specific data points (number of manufacturers, investment levels, job numbers), geographic positioning advantage, and how it creates sustained demand for the client. 3\u20134 sentences.]"
regional_num = 0
for p in doc.paragraphs:
    full = "".join(r.text for r in p.runs)
    if REGIONAL_DESC in full:
        regional_num += 1
        replace_para_text(p, REGIONAL_DESC, f"[Regional Factor {regional_num} Detail]")

# Also check tables for regional desc
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if REGIONAL_DESC in full_cell_text(cell):
                regional_num += 1
                replace_cell_text(cell, REGIONAL_DESC, f"[Regional Factor {regional_num} Detail]")

# Risk/mitigation table - 5 rows
risk_num = 0
RISK_DESC = "[Describe the market, competitive, operational, or macro risk.]"
RISK_MIT = "[Describe the specific action or positioning that reduces or eliminates this risk.]"
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        cells_text = [full_cell_text(c) for c in row.cells]
        if any(RISK_DESC in t for t in cells_text):
            risk_num += 1
            for ci, cell in enumerate(row.cells):
                replace_cell_text(cell, RISK_DESC, f"[Risk {risk_num}]")
                replace_cell_text(cell, RISK_MIT, f"[Risk {risk_num} Mitigation]")

# Milestone tables - [Key milestone / assumption for this year] x5, [Key milestone] x5
milestone_num = 0
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = full_cell_text(cell)
            if "[Key milestone / assumption for this year]" in txt:
                milestone_num += 1
                YEARS = [2026, 2027, 2028, 2029, 2030]
                yr = YEARS[milestone_num - 1] if milestone_num <= 5 else 2025 + milestone_num
                replace_cell_text(cell, "[Key milestone / assumption for this year]", f"[Forecast {yr} Assumption]")

# Timeline milestones
timeline_num = 0
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if full_cell_text(cell).strip() == "[Key milestone]":
                timeline_num += 1
                YEARS = [2026, 2027, 2028, 2029, 2030]
                yr = YEARS[timeline_num - 1] if timeline_num <= 5 else 2025 + timeline_num
                for p in cell.paragraphs:
                    if p.text.strip() == "[Key milestone]" and p.runs:
                        p.runs[0].text = f"[Timeline {yr} Milestone]"

# Location advantage
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            replace_cell_text(cell, "[One sentence \u2014 why this specific location is strategically positioned for outsized growth in this market.]",
                            "[Location Advantage]")

doc.save(os.path.join(TEMPLATE_DIR, "MKT-MF01-MWS-2026-04.06-Market Analysis 5-Year Forecast TEMPLATE.docx"))
print("  OK: Market Analysis")


# ═══════════════════════════════════════════════════════════════════════════════
# 5. REVENUE ACCELERATION PLAYBOOK - unique stat cards, guarantees, bonuses
# ═══════════════════════════════════════════════════════════════════════════════
print("Rebuilding: Revenue Acceleration Playbook...")
doc = Document(os.path.join(TEMPLATE_DIR, "MKT-PB03-MWS-2026-04.06-Revenue Acceleration Playbook TEMPLATE.docx"))

# Stat card table
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            replace_cell_text(cell, "[KEY CREDENTIAL]", "[Stat 2 Label]")
            replace_cell_text(cell, "[SCALE METRIC]", "[Stat 3 Label]")
            replace_cell_text(cell, "[CERT]", "[Stat 4 Cert]")
            replace_cell_text(cell, "[ABBREV]", "[Stat 4 Abbrev]")

# Bonus descriptions - numbered
bonus_num = 0
for p in doc.paragraphs:
    full = "".join(r.text for r in p.runs)
    for bn in range(1, 4):
        pat = f"[Describe bonus {bn}"
        if pat in full:
            # Replace the whole thing up to the next ]
            # These have unclosed brackets: [Describe bonus 1 — what it includes, anchored value ($[X,000]
            # Find the full text and replace
            replace_para_text(p, full.strip(), f"[Bonus {bn} Description]")

doc.save(os.path.join(TEMPLATE_DIR, "MKT-PB03-MWS-2026-04.06-Revenue Acceleration Playbook TEMPLATE.docx"))
print("  OK: Revenue Acceleration Playbook")


# ═══════════════════════════════════════════════════════════════════════════════
# 6. BLIND SPOT AUDIT - already mostly working, just fix the repeated [Fix] pattern
# ═══════════════════════════════════════════════════════════════════════════════
print("Rebuilding: Blind Spot Audit...")
doc = Document(os.path.join(TEMPLATE_DIR, "MKT-AU02-MWS-2026-04.06-Blind Spot Audit TEMPLATE.docx"))

# The fix placeholders repeat 5 times. Make them unique.
blind_spot_num = 0
for i, p in enumerate(doc.paragraphs):
    full = "".join(r.text for r in p.runs)

    # Detect blind spot section by the description paragraphs
    if "Describe the blind spot" in full:
        blind_spot_num = 1
    elif "buyer questions are going unanswered" in full:
        blind_spot_num = 2
    elif "digital channels the buyer uses" in full:
        blind_spot_num = 3
    elif "friction points in the buyer experience" in full:
        blind_spot_num = 4
    elif "single most powerful differentiator is buried" in full:
        blind_spot_num = 5

    if blind_spot_num > 0:
        # Replace fix placeholders with unique versions
        if "[Fix #1 \u2014 specific, actionable, one sentence]" in full:
            replace_para_text(p, "[Fix #1 \u2014 specific, actionable, one sentence]", f"[Fix {blind_spot_num}-1]")
        elif "[Fix #1]" in full:
            replace_para_text(p, "[Fix #1]", f"[Fix {blind_spot_num}-1]")
        if "[Fix #2]" in full:
            replace_para_text(p, "[Fix #2]", f"[Fix {blind_spot_num}-2]")
        if "[Fix #3]" in full:
            replace_para_text(p, "[Fix #3]", f"[Fix {blind_spot_num}-3]")
        if "[Fix #4]" in full:
            replace_para_text(p, "[Fix #4]", f"[Fix {blind_spot_num}-4]")

doc.save(os.path.join(TEMPLATE_DIR, "MKT-AU02-MWS-2026-04.06-Blind Spot Audit TEMPLATE.docx"))
print("  OK: Blind Spot Audit")

print("\nAll templates rebuilt with unique placeholders.")
print("Templates directory:", TEMPLATE_DIR)
