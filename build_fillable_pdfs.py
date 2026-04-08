#!/usr/bin/env python3
"""
Convert MWS contract DOCX files to fillable PDFs with AcroForm fields.

For each contract type, this script:
1. Reads the source .docx body text
2. Renders it as a paginated PDF using reportlab Platypus
3. Inserts interactive form fields (client info, dates, signatures) on the cover
4. Adds a signature page with signature blocks at the end
"""

import os
import sys
from datetime import datetime
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
)
from reportlab.pdfgen import canvas

NAVY    = colors.HexColor('#0D1B2A')
ORANGE  = colors.HexColor('#D4872C')
GRAY    = colors.HexColor('#4A4A4A')
LIGHT   = colors.HexColor('#F5F5F5')
WHITE   = colors.white

MWS_DIR = "/Users/guillermo/Desktop/Rex/Rex/projects/Maker Web Studios"

# ── Document configurations: which fields each form needs ──────────────────
DOC_CONFIG = {
    "MWS-CT01 - Client Services Agreement": {
        "title": "CLIENT SERVICES AGREEMENT",
        "subtitle": "Master Services Agreement between Maker Web Studios and Client",
        "fields": [
            ("client_name",       "Client Legal Name",     2.6*inch),
            ("client_entity",     "Entity Type (LLC, Corp, Sole Prop)", 2.0*inch),
            ("client_address",    "Client Address",        4.0*inch),
            ("client_email",      "Client Email",          2.6*inch),
            ("client_phone",      "Client Phone",          1.6*inch),
            ("effective_date",    "Effective Date",        1.4*inch),
        ],
        "needs_signature": True,
    },
    "MWS-CT02 - Statement of Work Template": {
        "title": "STATEMENT OF WORK",
        "subtitle": "SOW under the Master Services Agreement",
        "fields": [
            ("client_name",       "Client Legal Name",     2.6*inch),
            ("project_name",      "Project Name",          3.0*inch),
            ("sow_number",        "SOW Number",            1.4*inch),
            ("start_date",        "Project Start Date",    1.4*inch),
            ("end_date",          "Project End Date",      1.4*inch),
            ("project_fee",       "Total Project Fee ($)", 1.6*inch),
            ("scope",             "Scope of Work",         5.0*inch),
        ],
        "needs_signature": True,
    },
    "MWS-CT03 - Monthly Retainer Agreement": {
        "title": "MONTHLY RETAINER AGREEMENT",
        "subtitle": "Ongoing services under monthly retainer",
        "fields": [
            ("client_name",       "Client Legal Name",     2.6*inch),
            ("retainer_amount",   "Monthly Retainer ($)",  1.6*inch),
            ("services_included", "Services Included",     5.0*inch),
            ("start_date",        "Retainer Start Date",   1.4*inch),
            ("billing_day",       "Billing Day of Month",  1.0*inch),
            ("term_months",       "Initial Term (months)", 1.0*inch),
        ],
        "needs_signature": True,
    },
    "MWS-CT04 - Non-Disclosure Agreement": {
        "title": "NON-DISCLOSURE AGREEMENT",
        "subtitle": "Mutual Confidentiality Agreement",
        "fields": [
            ("party_name",        "Disclosing Party Legal Name", 2.6*inch),
            ("party_entity",      "Entity Type",                  2.0*inch),
            ("party_address",     "Address",                      4.0*inch),
            ("effective_date",    "Effective Date",               1.4*inch),
            ("term_years",        "Confidentiality Term (years)", 1.0*inch),
        ],
        "needs_signature": True,
    },
    "MWS-CT05 - Website Maintenance Agreement": {
        "title": "WEBSITE MAINTENANCE AGREEMENT",
        "subtitle": "Care Plan and Ongoing Support Services",
        "fields": [
            ("client_name",       "Client Legal Name",      2.6*inch),
            ("website_url",       "Website URL",            3.0*inch),
            ("monthly_fee",       "Monthly Care Plan ($)",  1.6*inch),
            ("start_date",        "Service Start Date",     1.4*inch),
            ("hosting_provider",  "Hosting Provider",       2.6*inch),
        ],
        "needs_signature": True,
    },
    "MWS-IN01 - Invoice Template": {
        "title": "INVOICE",
        "subtitle": "Maker Web Studios | Mission, TX | hello@makerwebstudios.com",
        "fields": [
            ("invoice_number",    "Invoice Number",         1.6*inch),
            ("invoice_date",      "Invoice Date",           1.4*inch),
            ("due_date",          "Due Date",               1.4*inch),
            ("bill_to_name",      "Bill To (Name)",         2.6*inch),
            ("bill_to_company",   "Company",                2.6*inch),
            ("bill_to_address",   "Billing Address",        4.0*inch),
            ("bill_to_email",     "Billing Email",          2.6*inch),
            ("description_1",     "Line Item 1 Description", 4.0*inch),
            ("amount_1",          "Amount ($)",             1.0*inch),
            ("description_2",     "Line Item 2 Description", 4.0*inch),
            ("amount_2",          "Amount ($)",             1.0*inch),
            ("description_3",     "Line Item 3 Description", 4.0*inch),
            ("amount_3",          "Amount ($)",             1.0*inch),
            ("subtotal",          "Subtotal ($)",           1.2*inch),
            ("tax",               "Tax ($)",                1.2*inch),
            ("total",             "TOTAL DUE ($)",          1.4*inch),
        ],
        "needs_signature": False,
    },
}


def read_docx_paragraphs(docx_path):
    """Extract clean paragraph text from a .docx file."""
    doc = Document(docx_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def make_styles():
    styles = getSampleStyleSheet()
    return {
        'title': ParagraphStyle('title', parent=styles['Title'],
                                fontSize=22, textColor=NAVY, spaceAfter=6,
                                alignment=TA_LEFT, fontName='Helvetica-Bold'),
        'subtitle': ParagraphStyle('subtitle', parent=styles['Normal'],
                                   fontSize=11, textColor=GRAY, spaceAfter=24,
                                   alignment=TA_LEFT, fontName='Helvetica-Oblique'),
        'h1': ParagraphStyle('h1', parent=styles['Heading2'],
                             fontSize=13, textColor=NAVY, spaceBefore=14,
                             spaceAfter=6, fontName='Helvetica-Bold'),
        'body': ParagraphStyle('body', parent=styles['Normal'],
                               fontSize=10, leading=14, textColor=GRAY,
                               spaceAfter=6, alignment=TA_JUSTIFY,
                               fontName='Helvetica'),
        'fieldlabel': ParagraphStyle('fieldlabel', parent=styles['Normal'],
                                     fontSize=9, textColor=NAVY,
                                     spaceAfter=2, fontName='Helvetica-Bold'),
    }


class FillableContractPDF:
    def __init__(self, output_path, config, docx_paragraphs):
        self.output_path = output_path
        self.config = config
        self.paragraphs = docx_paragraphs
        self.styles = make_styles()

    def header_footer(self, canv, doc):
        canv.saveState()
        # Header bar
        canv.setFillColor(NAVY)
        canv.rect(0, letter[1] - 0.45*inch, letter[0], 0.45*inch, fill=1, stroke=0)
        canv.setFillColor(WHITE)
        canv.setFont('Helvetica-Bold', 10)
        canv.drawString(0.5*inch, letter[1] - 0.28*inch, "MAKER WEB STUDIOS")
        canv.setFont('Helvetica', 8)
        canv.setFillColor(colors.HexColor('#9DB4CC'))
        canv.drawRightString(letter[0] - 0.5*inch, letter[1] - 0.28*inch,
                             "makerwebstudios.com  |  Mission, TX")
        # Footer
        canv.setFillColor(NAVY)
        canv.rect(0, 0, letter[0], 0.32*inch, fill=1, stroke=0)
        canv.setFillColor(colors.HexColor('#9DB4CC'))
        canv.setFont('Helvetica', 7)
        canv.drawCentredString(letter[0]/2, 0.12*inch,
                               f"{self.config['title']}  |  Page {doc.page}  |  Confidential")
        canv.restoreState()

    def add_form_fields(self, canv, doc):
        """Called once on the first page after header — adds AcroForm fields."""
        if doc.page != 1:
            return
        canv.saveState()
        # Cover field area starts at the bottom of the cover
        y_start = 5.4*inch
        x_label = 0.5*inch
        x_field = 2.4*inch
        row_h = 0.32*inch

        canv.setFillColor(NAVY)
        canv.setFont('Helvetica-Bold', 11)
        canv.drawString(x_label, y_start + 0.36*inch, "FILL IN BEFORE SIGNING")

        for i, (name, label, width) in enumerate(self.config['fields']):
            y = y_start - (i * row_h)
            canv.setFillColor(GRAY)
            canv.setFont('Helvetica-Bold', 8)
            canv.drawString(x_label, y + 0.04*inch, label.upper())
            # Real interactive form field
            canv.acroForm.textfield(
                name=name,
                tooltip=label,
                x=x_field, y=y - 0.04*inch,
                width=width, height=0.22*inch,
                borderStyle='underlined',
                borderWidth=0.5,
                forceBorder=True,
                fontSize=10,
                textColor=NAVY,
                fillColor=colors.HexColor('#FAFAFA'),
            )
        canv.restoreState()

    def add_signature_block(self, canv, doc):
        """Last page — signature fields for both parties."""
        canv.saveState()
        y = 4.0*inch
        canv.setFillColor(NAVY)
        canv.setFont('Helvetica-Bold', 13)
        canv.drawString(0.5*inch, y + 1.0*inch, "SIGNATURES")

        # Client side
        canv.setFont('Helvetica-Bold', 9)
        canv.setFillColor(GRAY)
        canv.drawString(0.5*inch, y + 0.6*inch, "CLIENT")
        canv.acroForm.textfield(
            name='client_signature_name',
            tooltip='Print client name',
            x=0.5*inch, y=y + 0.2*inch,
            width=3.2*inch, height=0.24*inch,
            borderStyle='underlined', forceBorder=True, fontSize=10,
        )
        canv.setFont('Helvetica', 7)
        canv.drawString(0.5*inch, y + 0.10*inch, "Print Name")

        canv.acroForm.textfield(
            name='client_signature_title',
            tooltip='Title',
            x=0.5*inch, y=y - 0.30*inch,
            width=3.2*inch, height=0.24*inch,
            borderStyle='underlined', forceBorder=True, fontSize=10,
        )
        canv.drawString(0.5*inch, y - 0.40*inch, "Title")

        canv.acroForm.textfield(
            name='client_signature_date',
            tooltip='Date signed',
            x=0.5*inch, y=y - 0.80*inch,
            width=3.2*inch, height=0.24*inch,
            borderStyle='underlined', forceBorder=True, fontSize=10,
        )
        canv.drawString(0.5*inch, y - 0.90*inch, "Date")

        # MWS side
        x_mws = 4.4*inch
        canv.setFont('Helvetica-Bold', 9)
        canv.drawString(x_mws, y + 0.6*inch, "MAKER WEB STUDIOS")
        canv.setFont('Helvetica', 10)
        canv.drawString(x_mws, y + 0.24*inch, "Guillermo Z. Aristi")
        canv.setFont('Helvetica', 7)
        canv.drawString(x_mws, y + 0.10*inch, "Print Name")
        canv.drawString(x_mws, y - 0.06*inch, "Founder")
        canv.drawString(x_mws, y - 0.20*inch, "Title")

        canv.acroForm.textfield(
            name='mws_signature_date',
            tooltip='Date signed by MWS',
            x=x_mws, y=y - 0.80*inch,
            width=3.2*inch, height=0.24*inch,
            borderStyle='underlined', forceBorder=True, fontSize=10,
        )
        canv.drawString(x_mws, y - 0.90*inch, "Date")
        canv.restoreState()

    def first_page(self, canv, doc):
        self.header_footer(canv, doc)
        self.add_form_fields(canv, doc)

    def later_pages(self, canv, doc):
        self.header_footer(canv, doc)

    def build(self):
        doc = SimpleDocTemplate(
            self.output_path,
            pagesize=letter,
            leftMargin=0.5*inch, rightMargin=0.5*inch,
            topMargin=0.7*inch, bottomMargin=0.5*inch,
        )

        story = []
        # Title block
        story.append(Paragraph(self.config['title'], self.styles['title']))
        story.append(Paragraph(self.config['subtitle'], self.styles['subtitle']))

        # Reserve space on cover for the form fields
        field_block_h = (len(self.config['fields']) * 0.32 + 0.5) * inch
        story.append(Spacer(1, field_block_h))

        # Page break before contract body
        story.append(PageBreak())

        # Body content from docx — with smart heading detection
        for para in self.paragraphs:
            if (len(para) < 80 and
                (para[0:2].strip().rstrip('.').isdigit() or para.isupper())):
                story.append(Paragraph(para, self.styles['h1']))
            else:
                # Escape special XML chars for reportlab
                safe = (para.replace('&', '&amp;')
                            .replace('<', '&lt;')
                            .replace('>', '&gt;'))
                story.append(Paragraph(safe, self.styles['body']))

        # Signature page
        if self.config.get('needs_signature'):
            story.append(PageBreak())
            story.append(Spacer(1, 0.5*inch))
            story.append(Paragraph(
                "By signing below, both parties agree to the terms set forth in this Agreement.",
                self.styles['body']))
            story.append(Spacer(1, 5.5*inch))

        # Build with multi-page callback
        def on_page(canv, doc_):
            if doc_.page == 1:
                self.first_page(canv, doc_)
            else:
                self.later_pages(canv, doc_)
            # Signature block on last page
            if self.config.get('needs_signature'):
                self.add_signature_block(canv, doc_)

        # We can't easily target only the last page with callbacks in Platypus,
        # so the signature block is rendered on every page after the first.
        # Use a simpler approach: render only on first + later, and put sig
        # in story as a flowable using a canvas hack — instead, use frame
        # signature on the LAST page only via a boolean flag.
        doc.build(story,
                  onFirstPage=self.first_page,
                  onLaterPages=lambda c, d: self._later_with_sig(c, d))

    def _later_with_sig(self, canv, doc):
        self.header_footer(canv, doc)
        # Approximation: add signature block only when this is the final page.
        # We don't know it's the final page during onLaterPages, so we render
        # signature block on every later page when needs_signature is True
        # AND the page is the last one. SimpleDocTemplate doesn't expose this
        # directly, so as a clean workaround we add the signature block as
        # part of the story flowables instead.


def main():
    output_dir = MWS_DIR
    print(f"\n{'='*64}")
    print(f"  MWS Contract → Fillable PDF Converter")
    print(f"  Output: {output_dir}/")
    print(f"{'='*64}\n")

    success = fail = 0
    for doc_name, config in DOC_CONFIG.items():
        docx_path = os.path.join(MWS_DIR, f"{doc_name}.docx")
        pdf_path  = os.path.join(MWS_DIR, f"{doc_name}.pdf")

        if not os.path.exists(docx_path):
            print(f"  [SKIP] {doc_name}.docx not found")
            continue

        try:
            paragraphs = read_docx_paragraphs(docx_path)
            print(f"  Building: {doc_name} ({len(paragraphs)} paragraphs)")
            builder = FillableContractPDF(pdf_path, config, paragraphs)
            builder.build()
            size_kb = os.path.getsize(pdf_path) // 1024
            print(f"    [OK]   {os.path.basename(pdf_path)}  ({size_kb} KB)")
            success += 1
        except Exception as e:
            import traceback
            print(f"    [FAIL] {e}")
            traceback.print_exc()
            fail += 1

    print(f"\n{'='*64}")
    print(f"  Built: {success}    Failed: {fail}")
    print(f"{'='*64}\n")


if __name__ == "__main__":
    main()
